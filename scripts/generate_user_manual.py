from __future__ import annotations

import math
import shutil
from datetime import date
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "manual"
BUILD_DIR = ROOT / ".build" / "user-manual"
ASSET_DIR = BUILD_DIR / "assets"
OUTPUT_DOCX = OUTPUT_DIR / "Gas_EC_Studio_RC6_User_Manual_ZH.docx"

APP_VERSION = "0.1.0 RC6"
MANUAL_VERSION = "1.0"
RELEASE_DATE = "2026-07-21"
PRODUCT_NAME = "Gas EC Studio"

FONT_CJK = "Microsoft YaHei"
FONT_LATIN = "Arial"
FONT_MONO = "Consolas"
FONT_PATH = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD_PATH = Path("C:/Windows/Fonts/msyhbd.ttc")

TEAL = "#0F5F63"
TEAL_DARK = "#0B3E42"
TEAL_LIGHT = "#E8F5F4"
CYAN = "#39A9AD"
AMBER = "#E7A23B"
AMBER_LIGHT = "#FFF5DF"
RED = "#C94C4C"
RED_LIGHT = "#FDECEC"
GREEN = "#2F7D59"
GREEN_LIGHT = "#EAF6EF"
INK = "#263238"
MUTED = "#607277"
LINE = "#CCDADB"
PALE = "#F5F8F8"
WHITE = "#FFFFFF"

PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_LEFT_MM = 17
MARGIN_RIGHT_MM = 17
MARGIN_TOP_MM = 18
MARGIN_BOTTOM_MM = 17
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(path), size=size)


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#"))


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def _set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_mm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Mm(width_mm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                value = str(edge_data[key])
                edge.set(qn(f"w:{key}"), value.lstrip("#") if key == "color" else value)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, mono: bool = False) -> None:
    font_name = FONT_MONO if mono else FONT_CJK
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_MONO if mono else FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_MONO if mono else FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = _rgb(color)


def _set_paragraph_keep(paragraph, keep_next: bool = False, keep_lines: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))
    p_pr.append(OxmlElement("w:widowControl"))


def _set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def _add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(text)
    run.append(end)
    paragraph._p.append(run)


def _set_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _draw_text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - width) / 2, (y0 + y1 - height) / 2), text, font=font, fill=fill, spacing=5, align="center")


def _save_flux_concept(path: Path) -> None:
    width, height = 1800, 920
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 55, width - 55, height - 55), radius=24, fill=PALE, outline=LINE, width=3)
    draw.text((105, 85), "涡动协方差：同相脉动决定通量方向与大小", font=_font(46, True), fill=TEAL_DARK)
    left, right = 120, 1670
    top, bottom = 245, 720
    mid1, mid2 = 390, 595
    draw.line((left, mid1, right, mid1), fill="#9AB0B3", width=3)
    draw.line((left, mid2, right, mid2), fill="#9AB0B3", width=3)
    t = np.linspace(0, 8 * math.pi, 680)
    w_prime = 0.72 * np.sin(t) + 0.23 * np.sin(2.7 * t + 0.4)
    c_prime = 0.86 * w_prime + 0.22 * np.sin(1.4 * t - 0.6)
    xs = np.linspace(left, right, t.size)
    w_pts = [(float(x), float(mid1 - y * 105)) for x, y in zip(xs, w_prime)]
    c_pts = [(float(x), float(mid2 - y * 105)) for x, y in zip(xs, c_prime)]
    draw.line(w_pts, fill=TEAL, width=6, joint="curve")
    draw.line(c_pts, fill=AMBER, width=6, joint="curve")
    draw.text((120, 205), "垂直风速脉动 w'", font=_font(34, True), fill=TEAL)
    draw.text((120, 450), "标量脉动 c'", font=_font(34, True), fill=AMBER)
    for idx in (75, 235, 395, 555):
        x = int(xs[idx])
        sign = 1 if w_prime[idx] * c_prime[idx] > 0 else -1
        color = GREEN if sign > 0 else RED
        draw.line((x, mid1 + 118, x, mid2 - 118), fill=color, width=3)
    draw.rounded_rectangle((1060, 755, 1660, 850), radius=16, fill=TEAL_LIGHT, outline=CYAN, width=2)
    _draw_text_center(draw, (1060, 755, 1660, 850), "F = 平均(w'c')；同号占优为正，异号占优为负", _font(28, True), TEAL_DARK)
    draw.text((120, 785), "先分解平均量与脉动量，再在同一平均窗口内求协方差。", font=_font(28), fill=MUTED)
    image.save(path, quality=95)


def _save_processing_chain(path: Path) -> None:
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "从高频样本到可审计通量的处理链", font=_font(48, True), fill=TEAL_DARK)
    items = [
        ("1  输入与时间", "字段映射、单位、时钟、采样率、缺测"),
        ("2  原始质量", "范围、尖峰、幅值、掉线、诊断字"),
        ("3  统计预处理", "平均窗口、去趋势、坐标旋转、时滞"),
        ("4  通量计算", "协方差、热力学量、密度修正"),
        ("5  频率响应", "功率谱、协谱、Ogive、传递函数"),
        ("6  科学评估", "稳态、湍流、u*、不确定度、足迹"),
        ("7  结果交付", "QC、方法快照、报告、证据与清单"),
    ]
    colors = [TEAL, "#267F84", "#2F9699", AMBER, "#D8882D", "#557A7D", TEAL_DARK]
    box_w, box_h = 455, 180
    positions = [(80, 180), (670, 180), (1260, 180), (1260, 480), (670, 480), (80, 480), (670, 780)]
    for index, ((title, note), (x, y), color) in enumerate(zip(items, positions, colors)):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=WHITE, outline=color, width=5)
        draw.rounded_rectangle((x, y, x + box_w, y + 56), radius=14, fill=color)
        draw.rectangle((x, y + 38, x + box_w, y + 56), fill=color)
        draw.text((x + 24, y + 12), title, font=_font(30, True), fill=WHITE)
        _draw_text_center(draw, (x + 24, y + 66, x + box_w - 24, y + box_h - 12), note, _font(26), INK)
        if index < len(items) - 1:
            x2, y2 = positions[index + 1]
            start_x = x + box_w if x2 > x else x if x2 < x else x + box_w // 2
            end_x = x2 if x2 > x else x2 + box_w if x2 < x else x2 + box_w // 2
            start_y = y + box_h // 2 if y2 == y else y + box_h
            end_y = y2 + box_h // 2 if y2 == y else y2
            draw.line((start_x, start_y, end_x, end_y), fill="#91A8AA", width=8)
            angle = math.atan2(end_y - start_y, end_x - start_x)
            p1 = (end_x - 22 * math.cos(angle - 0.5), end_y - 22 * math.sin(angle - 0.5))
            p2 = (end_x - 22 * math.cos(angle + 0.5), end_y - 22 * math.sin(angle + 0.5))
            draw.polygon([(end_x, end_y), p1, p2], fill="#91A8AA")
    image.save(path, quality=95)


def _save_lag_curve(path: Path) -> None:
    width, height = 1700, 900
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "时滞补偿：在物理合理窗口内寻找协方差峰值", font=_font(46, True), fill=TEAL_DARK)
    x0, y0, x1, y1 = 160, 170, 1570, 720
    draw.line((x0, y1, x1, y1), fill=INK, width=4)
    draw.line((x0, y0, x0, y1), fill=INK, width=4)
    lags = np.linspace(-2.0, 3.0, 700)
    cov = 0.78 * np.exp(-((lags - 0.55) / 0.55) ** 2) - 0.12 * np.exp(-((lags + 0.8) / 0.7) ** 2)
    xs = x0 + (lags + 2.0) / 5.0 * (x1 - x0)
    ys = y1 - (cov + 0.2) / 1.1 * (y1 - y0)
    draw.line([(float(x), float(y)) for x, y in zip(xs, ys)], fill=TEAL, width=7, joint="curve")
    peak_idx = int(np.argmax(cov))
    px, py = int(xs[peak_idx]), int(ys[peak_idx])
    draw.line((px, py, px, y1), fill=AMBER, width=4)
    draw.ellipse((px - 11, py - 11, px + 11, py + 11), fill=AMBER)
    draw.rounded_rectangle((px + 35, py - 70, px + 420, py + 20), radius=14, fill=AMBER_LIGHT, outline=AMBER, width=2)
    _draw_text_center(draw, (px + 35, py - 70, px + 420, py + 20), "峰值时滞约 0.55 s", _font(29, True), INK)
    for lag in range(-2, 4):
        x = int(x0 + (lag + 2.0) / 5.0 * (x1 - x0))
        draw.line((x, y1, x, y1 + 12), fill=INK, width=3)
        draw.text((x - 16, y1 + 20), str(lag), font=_font(24), fill=MUTED)
    draw.text((740, 795), "时滞 / s", font=_font(30, True), fill=INK)
    draw.text((36, 395), "协方差", font=_font(30, True), fill=INK)
    draw.text((120, 835), "搜索窗应覆盖传感器间距、流速与管路输送造成的延迟；边界峰值应触发回退或复核。", font=_font(27), fill=MUTED)
    image.save(path, quality=95)


def _save_spectral_transfer(path: Path) -> None:
    width, height = 1750, 930
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "频率响应：观测协谱与理想协谱之间的衰减", font=_font(46, True), fill=TEAL_DARK)
    x0, y0, x1, y1 = 160, 170, 1600, 720
    draw.line((x0, y1, x1, y1), fill=INK, width=4)
    draw.line((x0, y0, x0, y1), fill=INK, width=4)
    f = np.logspace(-3, 1, 800)
    ideal = (f / 0.07) ** 0.7 / (1.0 + (f / 0.07) ** 1.65)
    transfer = 1.0 / (1.0 + (f / 0.85) ** 2.2)
    observed = ideal * transfer
    logx = np.log10(f)
    xs = x0 + (logx + 3.0) / 4.0 * (x1 - x0)
    ideal_y = y1 - ideal / ideal.max() * 455
    observed_y = y1 - observed / ideal.max() * 455
    draw.line([(float(x), float(y)) for x, y in zip(xs, ideal_y)], fill=AMBER, width=6)
    draw.line([(float(x), float(y)) for x, y in zip(xs, observed_y)], fill=TEAL, width=6)
    for decade, label in zip((-3, -2, -1, 0, 1), ("0.001", "0.01", "0.1", "1", "10")):
        x = int(x0 + (decade + 3.0) / 4.0 * (x1 - x0))
        draw.line((x, y1, x, y1 + 12), fill=INK, width=3)
        draw.text((x - 30, y1 + 22), label, font=_font(23), fill=MUTED)
    draw.line((1050, 200, 1140, 200), fill=AMBER, width=7)
    draw.text((1160, 180), "理想协谱", font=_font(27, True), fill=INK)
    draw.line((1050, 255, 1140, 255), fill=TEAL, width=7)
    draw.text((1160, 235), "观测协谱", font=_font(27, True), fill=INK)
    draw.text((680, 805), "频率 / Hz（对数轴）", font=_font(30, True), fill=INK)
    draw.text((45, 410), "相对贡献", font=_font(28, True), fill=INK)
    draw.text((115, 855), "高频端差异反映路径平均、传感器响应、空间分离、管路与数字滤波等影响。", font=_font(27), fill=MUTED)
    image.save(path, quality=95)


def _save_footprint(path: Path) -> None:
    width, height = 1750, 950
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "通量足迹：当前半小时结果主要来自哪里", font=_font(46, True), fill=TEAL_DARK)
    tower_x, tower_y = 260, 470
    draw.line((tower_x, 750, tower_x, 270), fill=INK, width=12)
    draw.polygon([(tower_x - 42, 300), (tower_x + 42, 300), (tower_x, 230)], fill=TEAL)
    draw.text((165, 770), "观测塔", font=_font(28, True), fill=INK)
    draw.line((365, 210, 700, 210), fill=AMBER, width=10)
    draw.polygon([(700, 210), (660, 185), (660, 235)], fill=AMBER)
    draw.text((390, 145), "平均风向", font=_font(27, True), fill=INK)
    bands = [
        (480, 170, "90%", "#D8EEEC"),
        (390, 135, "70%", "#B8DEDB"),
        (300, 105, "50%", "#82C5C1"),
        (205, 75, "30%", "#46A6A7"),
    ]
    peak_x = 590
    for length, half_height, label, fill in bands:
        box = (tower_x + 90, tower_y - half_height, tower_x + 90 + length, tower_y + half_height)
        draw.ellipse(box, fill=fill, outline=TEAL, width=3)
        draw.text((box[2] - 60, box[1] + 8), label, font=_font(25, True), fill=TEAL_DARK)
    draw.ellipse((peak_x - 12, tower_y - 12, peak_x + 12, tower_y + 12), fill=RED)
    draw.line((peak_x, tower_y - 15, peak_x, tower_y - 180), fill=RED, width=3)
    draw.rounded_rectangle((peak_x - 145, tower_y - 250, peak_x + 215, tower_y - 180), radius=12, fill=RED_LIGHT, outline=RED, width=2)
    _draw_text_center(draw, (peak_x - 145, tower_y - 250, peak_x + 215, tower_y - 180), "峰值贡献距离", _font(27, True), RED)
    draw.rounded_rectangle((980, 250, 1640, 705), radius=18, fill=PALE, outline=LINE, width=3)
    notes = [
        "输入：测量高度、粗糙度、冠层高度",
        "输入：u*、稳定度、风向与平均风速",
        "输出：峰值距离与累计贡献距离",
        "2D 输出可与地类栅格叠加核验",
        "低湍流、复杂地形与非均匀下垫面需谨慎",
    ]
    draw.text((1030, 285), "解释要点", font=_font(34, True), fill=TEAL_DARK)
    for i, note in enumerate(notes):
        y = 360 + i * 62
        draw.ellipse((1030, y + 8, 1046, y + 24), fill=AMBER)
        draw.text((1065, y), note, font=_font(25), fill=INK)
    draw.text((120, 860), "足迹不是通量修正因子；它回答来源区域与代表性问题。", font=_font(30, True), fill=MUTED)
    image.save(path, quality=95)


def create_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "flux": ASSET_DIR / "flux_concept.png",
        "chain": ASSET_DIR / "processing_chain.png",
        "lag": ASSET_DIR / "lag_curve.png",
        "spectral": ASSET_DIR / "spectral_transfer.png",
        "footprint": ASSET_DIR / "footprint.png",
    }
    _save_flux_concept(paths["flux"])
    _save_processing_chain(paths["chain"])
    _save_lag_curve(paths["lag"])
    _save_spectral_transfer(paths["spectral"])
    _save_footprint(paths["footprint"])
    return paths


class ManualBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self.figure_number = 0
        self.table_number = 0
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Mm(PAGE_WIDTH_MM)
        section.page_height = Mm(PAGE_HEIGHT_MM)
        section.top_margin = Mm(MARGIN_TOP_MM)
        section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
        section.left_margin = Mm(MARGIN_LEFT_MM)
        section.right_margin = Mm(MARGIN_RIGHT_MM)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = FONT_CJK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = _rgb(INK)
        normal.paragraph_format.line_spacing = 1.25
        normal.paragraph_format.space_after = Pt(5)

        for name, size, color, before, after in (
            ("Title", 30, WHITE, 0, 0),
            ("Subtitle", 13, WHITE, 0, 0),
            ("Heading 1", 21, TEAL_DARK, 0, 12),
            ("Heading 2", 14, TEAL, 13, 6),
            ("Heading 3", 11, INK, 10, 4),
            ("Caption", 8.5, MUTED, 4, 8),
        ):
            style = self.doc.styles[name]
            style.font.name = FONT_CJK
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
            style.font.size = Pt(size)
            style.font.color.rgb = _rgb(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        self.doc.styles["Heading 1"].font.bold = True
        self.doc.styles["Heading 1"].paragraph_format.page_break_before = True
        self.doc.styles["Heading 2"].font.bold = True
        self.doc.styles["Heading 3"].font.bold = True
        self.doc.styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for style_name in ("List Bullet", "List Number"):
            style = self.doc.styles[style_name]
            style.font.name = FONT_CJK
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
            style.font.size = Pt(9.3)
            style.paragraph_format.left_indent = Mm(6)
            style.paragraph_format.first_line_indent = Mm(-3)
            style.paragraph_format.space_after = Pt(3)

        self._configure_header_footer(section)
        props = self.doc.core_properties
        props.title = "Gas EC Studio 详细使用说明书"
        props.subject = "涡动协方差采集、处理、谱质控与交付"
        props.author = "Gas EC Studio 项目组"
        props.keywords = "涡动协方差, EC, 通量, 谱分析, 质量控制, Gas EC Studio"
        props.comments = f"软件版本 {APP_VERSION}; 手册版本 {MANUAL_VERSION}"

    def _configure_header_footer(self, section) -> None:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{PRODUCT_NAME}  |  详细使用说明书  |  {APP_VERSION}")
        _set_run_font(run, 7.5, color=MUTED)
        bottom = {"val": "single", "sz": "6", "space": "1", "color": LINE}
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        for key, value in bottom.items():
            b.set(qn(f"w:{key}"), value)
        p_bdr.append(b)
        p_pr.append(p_bdr)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(2)
        run = fp.add_run("第 ")
        _set_run_font(run, 7.5, color=MUTED)
        _add_field(fp, "PAGE", "1")
        run = fp.add_run(" 页  |  科学结果须结合站点元数据与 QC 证据解释")
        _set_run_font(run, 7.5, color=MUTED)

    def paragraph(self, text: str = "", *, bold_lead: str | None = None, align=None, space_after: float | None = None):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        if bold_lead and text.startswith(bold_lead):
            lead = p.add_run(bold_lead)
            _set_run_font(lead, bold=True, color=TEAL_DARK)
            body = p.add_run(text[len(bold_lead):])
            _set_run_font(body)
        else:
            run = p.add_run(text)
            _set_run_font(run)
        _set_paragraph_keep(p)
        return p

    def bullets(self, items: list[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            _set_run_font(run)
            _set_paragraph_keep(p)

    def steps(self, items: list[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph(style="List Number")
            run = p.add_run(item)
            _set_run_font(run)
            _set_paragraph_keep(p)

    def heading(self, text: str, level: int = 1):
        p = self.doc.add_heading(text, level=level)
        _set_paragraph_keep(p, keep_next=True)
        return p

    def callout(self, title: str, body: str, tone: str = "info") -> None:
        colors = {
            "info": (TEAL_LIGHT, TEAL, TEAL_DARK),
            "warning": (AMBER_LIGHT, AMBER, INK),
            "danger": (RED_LIGHT, RED, RED),
            "success": (GREEN_LIGHT, GREEN, GREEN),
        }
        fill, edge, text_color = colors[tone]
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_fixed_layout(table)
        _set_cell_width(table.cell(0, 0), 5)
        _set_cell_width(table.cell(0, 1), CONTENT_WIDTH_MM - 5)
        marker, content = table.rows[0].cells
        _set_cell_shading(marker, edge)
        _set_cell_shading(content, fill)
        _set_cell_margins(marker, top=80, start=0, bottom=80, end=0)
        _set_cell_margins(content, top=110, start=150, bottom=110, end=150)
        for cell in (marker, content):
            _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        p = content.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lead = p.add_run(title + "  ")
        _set_run_font(lead, 9.5, bold=True, color=text_color)
        run = p.add_run(body)
        _set_run_font(run, 9.2, color=INK)
        _prevent_row_split(table.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def formula(self, label: str, formula: str, explanation: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_fixed_layout(table)
        cell = table.cell(0, 0)
        _set_cell_width(cell, CONTENT_WIDTH_MM)
        _set_cell_shading(cell, PALE)
        _set_cell_margins(cell, top=140, start=190, bottom=140, end=190)
        _set_cell_border(cell, top={"val": "single", "sz": "8", "color": CYAN}, bottom={"val": "single", "sz": "8", "color": CYAN}, start={"val": "single", "sz": "8", "color": CYAN}, end={"val": "single", "sz": "8", "color": CYAN})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        lead = p.add_run(label + "  ")
        _set_run_font(lead, 9.5, bold=True, color=TEAL)
        value = p.add_run(formula)
        _set_run_font(value, 12.5, bold=True, color=INK)
        q = cell.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(0)
        run = q.add_run(explanation)
        _set_run_font(run, 8.6, color=MUTED)
        _prevent_row_split(table.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def table(self, headers: list[str], rows: list[list[str]], widths_mm: list[float], caption: str | None = None) -> None:
        if len(headers) != len(widths_mm):
            raise ValueError("table headers and widths do not match")
        if abs(sum(widths_mm) - CONTENT_WIDTH_MM) > 1.5:
            raise ValueError(f"table widths sum to {sum(widths_mm)}, expected {CONTENT_WIDTH_MM}")
        self.table_number += 1
        if caption:
            p = self.doc.add_paragraph(style="Caption")
            run = p.add_run(f"表 {self.table_number}  {caption}")
            _set_run_font(run, 8.5, color=MUTED)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_fixed_layout(table)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, (cell, text, width) in enumerate(zip(header_cells, headers, widths_mm)):
            _set_cell_width(cell, width)
            _set_cell_shading(cell, TEAL_DARK)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            _set_run_font(run, 8.4, bold=True, color=WHITE)
        _set_repeat_table_header(table.rows[0])
        _prevent_row_split(table.rows[0])
        for row_index, row_data in enumerate(rows):
            row = table.add_row()
            _prevent_row_split(row)
            for cell, text, width in zip(row.cells, row_data, widths_mm):
                _set_cell_width(cell, width)
                _set_cell_margins(cell)
                _set_cell_shading(cell, WHITE if row_index % 2 == 0 else PALE)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(text))
                _set_run_font(run, 8.2, color=INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def figure(self, path: Path, caption: str, *, width_mm: float = CONTENT_WIDTH_MM, alt: str | None = None) -> None:
        if not path.exists():
            return
        self.figure_number += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        inline = p.add_run().add_picture(str(path), width=Mm(width_mm))
        _set_alt_text(inline, f"图 {self.figure_number}", alt or caption)
        cp = self.doc.add_paragraph(style="Caption")
        run = cp.add_run(f"图 {self.figure_number}  {caption}")
        _set_run_font(run, 8.5, color=MUTED)
        _set_paragraph_keep(cp)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def add_toc(self) -> None:
        self.heading("目录", 1)
        p = self.doc.add_paragraph()
        _add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "在 Word 中更新目录域")
        self.callout("目录更新", "DOCX 首次打开时如页码未刷新，请在目录内右键选择“更新域 -> 更新整个目录”。发布 PDF 已在导出前刷新目录。", "info")

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)


def screenshot(name: str) -> Path:
    return ROOT / ".build" / "deep-audit" / "ui-current" / name


def add_cover(builder: ManualBuilder) -> None:
    section = builder.doc.sections[0]
    section.different_first_page_header_footer = True
    cover = builder.doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    _set_fixed_layout(cover)
    cell = cover.cell(0, 0)
    _set_cell_width(cell, CONTENT_WIDTH_MM)
    _set_cell_shading(cell, TEAL_DARK)
    _set_cell_margins(cell, top=360, start=360, bottom=330, end=360)
    _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
    logo_path = ROOT / "app" / "assets" / "gas_ec_studio_icon.png"
    if logo_path.exists():
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        inline = p.add_run().add_picture(str(logo_path), width=Mm(24))
        _set_alt_text(inline, "Gas EC Studio 图标", "青绿色涡动通量软件图标")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(PRODUCT_NAME)
    _set_run_font(r, 30, bold=True, color=WHITE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("详细使用说明书")
    _set_run_font(r, 25, bold=True, color=WHITE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("涡动协方差采集 · 处理 · 谱质控 · 科学解释 · 证据交付")
    _set_run_font(r, 12.5, color="D7EFEF")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"软件版本  {APP_VERSION}     手册版本  {MANUAL_VERSION}     发布日期  {RELEASE_DATE}")
    _set_run_font(r, 9.5, bold=True, color="F7D79A")

    builder.doc.add_paragraph().paragraph_format.space_after = Pt(1)
    if screenshot("05-page-ec_processing.png").exists():
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        inline = p.add_run().add_picture(str(screenshot("05-page-ec_processing.png")), width=Mm(CONTENT_WIDTH_MM))
        _set_alt_text(inline, "EC 处理工作台", "Gas EC Studio 的 EC 处理页面全景；其中数值为演示数据")
    builder.callout("适用范围", "本手册面向现场操作员、通量工程师、科研人员与交付审核人员。界面截图中的数值均为演示数据，不构成真实站点科学结论。", "warning")
    p = builder.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Gas EC Studio 项目组")
    _set_run_font(r, 10, bold=True, color=TEAL_DARK)
    builder.page_break()


def add_document_control(builder: ManualBuilder) -> None:
    builder.heading("文档控制与阅读说明", 1)
    builder.table(
        ["项目", "内容"],
        [
            ["文档名称", "Gas EC Studio 详细使用说明书"],
            ["对应软件", APP_VERSION],
            ["手册版本", MANUAL_VERSION],
            ["发布日期", RELEASE_DATE],
            ["适用系统", "Windows 10/11 64 位；以发布包的系统要求为准"],
            ["默认数据目录", "%LOCALAPPDATA%\\GasECStudio\\runtime_data"],
            ["公开下载", "https://github.com/xuxin2023/gas_ec_studio/releases"],
            ["文档状态", "正式发布版；科学解释需结合项目元数据、QC 与证据包"],
        ],
        [40, 136],
        "文档识别信息",
    )
    builder.heading("如何使用本手册", 2)
    builder.table(
        ["角色", "优先阅读", "目标"],
        [
            ["现场操作员", "第 1、4、5、10、11 章", "接入设备、确认采集连续、完成日常检查与故障处置"],
            ["通量工程师", "第 2、6、7、8、12 章", "建立科学处理配置、解释谱与 QC、复核异常窗口"],
            ["项目负责人", "第 3、9、10、12 章", "确认项目元数据、方法版本、交付门槛与可追溯性"],
            ["科研审核人员", "第 2、6、7、8、附录 A-E", "核对假设、算法选择、限制条件与参考依据"],
        ],
        [29, 58, 89],
        "角色导读",
    )
    builder.callout("科学边界", "软件可以自动执行计算、诊断和证据整理，但不会替代研究者对下垫面均匀性、站点代表性、仪器校准、能量闭合与生态过程的科学判断。", "danger")
    builder.heading("符号与提示", 2)
    builder.table(
        ["标识", "含义", "处理原则"],
        [
            ["提示", "有助于提高效率或可重复性的建议", "优先采用，但可按项目 SOP 调整"],
            ["注意", "可能影响质量、完整性或可解释性的条件", "处理前核对元数据与证据"],
            ["警告", "可能导致错误通量、错误符号或不可交付结果", "停止交付并闭合问题"],
            ["演示数据", "截图或示例值用于说明界面", "不得用于真实科学报告"],
        ],
        [22, 72, 82],
        "提示级别",
    )


def add_chapter_1(builder: ManualBuilder) -> None:
    builder.heading("1  快速开始", 1)
    builder.paragraph("本章给出从安装到交付的最短闭环。首次使用建议完整走一遍；日常运行可直接使用第 10 章的 SOP 清单。")
    builder.heading("1.1  安装、启动与版本确认", 2)
    builder.steps([
        "从项目公开发布页下载与当前系统匹配的 Windows 发布包；不要从来源不明的镜像获取可执行文件。",
        "将便携包放在具有写权限的本地目录。正式项目不建议直接从压缩包内运行。",
        "启动后打开“关于”窗口，确认软件版本为 0.1.0 RC6，并记录发布批次。",
        "确认默认工作区可写，且磁盘空间能够容纳高频原始数据、派生结果、谱文件和证据包。",
        "正式采集前检查 Windows 时间、时区、休眠策略与安全软件拦截情况。",
    ])
    builder.figure(screenshot("48-about-version.png"), "关于窗口与版本信息。截图为软件状态示例。", width_mm=112, alt="关于窗口显示 Gas EC Studio 0.1.0 RC6")
    builder.callout("版本原则", "处理结果必须与软件版本、项目配置快照和输入批次绑定。方法参数发生变化时，应生成新批次而不是覆盖旧结果。", "warning")

    builder.heading("1.2  六个工作区的标准顺序", 2)
    builder.table(
        ["顺序", "工作区", "必须完成的动作", "通过标志"],
        [
            ["1", "设备中心", "连接、读回、确认在线/采集/告警", "设备状态与配置一致"],
            ["2", "实时采集", "检查曲线、缓存、时间戳和连续性", "高频流稳定且无持续缺口"],
            ["3", "项目与站点", "填写站点、仪器、布局、采样链和时间", "关键元数据无空项"],
            ["4", "EC 处理", "选择数据源、时间范围和处理方法", "窗口计算完成且有方法快照"],
            ["5", "谱修正与 QC", "复核谱形、时滞、Ogive、QC 和异常", "异常已解释或排除"],
            ["6", "报告中心", "生成报告、证据包、清单并检查门槛", "交付链与验证链闭合"],
        ],
        [13, 30, 83, 50],
        "标准工作流",
    )
    builder.heading("1.3  十分钟首轮运行", 2)
    builder.steps([
        "在设备中心选择接口并连接设备；仅在读回值与项目预期一致后开始采集。",
        "在实时采集页观察至少一个稳定时段，确认 CO2、H2O、风速、温度和压力的时间戳同步。",
        "在项目与站点页建立项目、站点、仪器布局和采样链；保存测量高度、传感器间距及管路参数。",
        "进入 EC 处理页，选择当前项目高频目录和合适的时间范围；初次建议使用 30 min 窗、双旋转、块均值、协方差最大带默认与密度修正。",
        "运行预检，先解决字段、单位、连续性和元数据问题，再启动正式处理。",
        "进入谱修正与 QC，检查时滞峰是否位于合理区间、协谱是否有物理形态、Ogive 是否收敛。",
        "在报告中心生成报告；查看异常事件、方法来源、证据清单和交付门槛。",
        "导出报告与证据包，在另一目录中验证文件可打开、清单完整、版本一致。",
    ])
    builder.callout("不要跳过预检", "时钟、单位、风向定义、测量高度或传感器几何错误，会在后续形成看似平滑但物理含义错误的通量。", "danger")
    builder.heading("1.4  OP 与 ENG 视图", 2)
    builder.paragraph("OP 视图面向日常操作，强调状态、下一步和闭环；ENG 视图显示方法参数、诊断细节、来源证据和适用边界。切换视图不会改变数据本身，但在保存配置前应确认当前可见参数。")


def add_chapter_2(builder: ManualBuilder, diagrams: dict[str, Path]) -> None:
    builder.heading("2  涡动协方差科学基础", 1)
    builder.paragraph("涡动协方差法利用高频垂直风速与气体、温度等标量的同步脉动，估计地表与大气之间的净交换。它不是简单的浓度平均，而是对共同脉动的统计测量。")
    builder.heading("2.1  Reynolds 分解与协方差", 2)
    builder.formula("Reynolds 分解", "x(t) = x̄ + x′(t)", "x̄ 为平均窗口内的均值，x′ 为瞬时值相对均值或趋势的脉动。")
    builder.formula("标量通量核心", "F_c = ρ · overline(w′c′)", "w′ 为垂直风速脉动，c′ 为标量脉动；ρ 与单位换算取决于标量表达形式。")
    builder.figure(diagrams["flux"], "垂直风速与标量同相脉动产生净通量的示意。", alt="两条同相脉动曲线说明协方差通量")
    builder.paragraph("当向上的 w′ 与高于平均值的 c′ 同时出现，或向下的 w′ 与低于平均值的 c′ 同时出现，乘积通常为正；相反组合使乘积为负。最终符号还取决于坐标轴方向、仪器字段定义与报告约定，因此项目必须固定“向上为正”的坐标约定。")
    builder.table(
        ["量", "典型表达", "常用单位", "解释"],
        [
            ["CO2 通量", "Fc", "µmol m⁻² s⁻¹", "正值通常表示向大气释放，负值通常表示下垫面吸收"],
            ["水汽通量", "E", "mmol m⁻² s⁻¹", "与潜热通量 LE 可相互转换"],
            ["显热通量", "H", "W m⁻²", "由 w′T′ 与空气热容量、密度确定"],
            ["潜热通量", "LE", "W m⁻²", "由水汽通量乘汽化潜热得到"],
            ["动量通量", "τ", "N m⁻²", "由风速协方差与空气密度确定"],
            ["摩擦速度", "u*", "m s⁻¹", "表征近地层湍流交换强度"],
        ],
        [30, 30, 40, 76],
        "主要通量与湍流量",
    )

    builder.heading("2.2  平均窗口为什么重要", 2)
    builder.paragraph("平均窗口需要在两个目标之间平衡：足够长，以采样主要湍涡并获得稳定协方差；又不能长到跨越天气突变、风向转换或非稳态过程。陆地生态通量常使用 30 min，但这不是不可更改的常数。短窗口适合诊断和快速过程，长窗口可能改善低频采样，却会增加非稳态风险。")
    builder.bullets([
        "窗口必须与采样率共同解释；10 Hz 的 30 min 理论样本数为 18 000。",
        "缺测率、有效样本数与连续性应随每个窗口保存。",
        "窗口边界应使用明确时区和时间戳归属规则，避免日/月汇总错位。",
        "改变窗口长度会改变去趋势、时滞、谱修正与稳态检验的含义，必须形成新方法版本。",
    ])

    builder.heading("2.3  成立条件与常见违背", 2)
    builder.table(
        ["科学条件", "理想状态", "常见违背", "软件中的证据"],
        [
            ["湍流充分", "交换由湍流主导", "夜间弱风、稳定层结、间歇湍流", "u*、湍流分数、稳定度、QC"],
            ["近似稳态", "窗口内统计性质变化小", "锋面、降雨、风向突变、启停", "六子窗稳态对比与异常事件"],
            ["水平平流可忽略", "垂直协方差代表主要交换", "坡地、林缘、城市峡谷", "风向、足迹、地类叠加与限制说明"],
            ["频率响应充分", "主要涡尺度被同步采样", "管路衰减、传感器分离、响应慢", "功率谱、协谱、传递函数、修正因子"],
            ["标量守恒", "窗口尺度内源汇可解释", "化学反应、存储项、水平输送", "项目方法说明与外部辅助数据"],
            ["元数据准确", "几何、单位、时钟和校准可信", "高度、间距、管长或时区填错", "预检、配置快照、方法来源"],
        ],
        [32, 42, 50, 52],
        "涡动协方差的核心假设",
    )
    builder.callout("QC 的含义", "质量标志用于说明某个窗口对既定科学用途是否可信，不会把不满足假设的数据自动变成满足假设的数据。", "warning")

    builder.heading("2.4  从原始样本到结果", 2)
    builder.figure(diagrams["chain"], "软件处理链及其审计节点。", alt="七阶段涡动协方差处理流程图")
    builder.paragraph("每一步都会改变后续统计量。正确的做法是保存方法选择、参数、输入批次和诊断，而不是只保留最后一列通量。Gas EC Studio 的报告中心和证据包正是为这一可追溯链设计。")


def add_chapter_3(builder: ManualBuilder) -> None:
    builder.heading("3  软件界面与工作区", 1)
    builder.paragraph("主窗口按任务链分成六个工作区。顶部显示全局状态与视图模式，左侧负责导航，中部是当前任务工作台，右侧提供上下文诊断或交付状态，底部运行日志用于定位错误。")
    builder.table(
        ["区域", "用途", "使用建议"],
        [
            ["顶部状态带", "版本、项目、采集、处理与交付状态", "切换页面前先确认当前项目和批次"],
            ["左侧导航", "进入六个主工作区及步骤目录", "按任务顺序使用，不以页面是否可点击替代质量闭环"],
            ["中央工作台", "配置、曲线、表格、方法和结果", "修改后先保存，再运行预检或处理"],
            ["右侧检查器", "显示当前步骤、参数来源、风险和下一步", "ENG 视图下重点复核适用边界"],
            ["底部日志", "记录连接、采集、处理、导出和错误", "异常时复制完整时间段，不只截最后一行"],
        ],
        [32, 72, 72],
        "主窗口区域",
    )
    builder.figure(screenshot("01-page-device_center.png"), "设备中心全景。界面数值为演示数据。", alt="设备中心包含连接状态、命令区和诊断信息")
    builder.heading("3.1  页面间的数据关系", 2)
    builder.paragraph("设备中心与实时采集形成“现场数据入口”；项目与站点形成“科学元数据入口”；EC 处理与谱修正形成“方法与结果入口”；报告中心形成“审阅与交付入口”。任何一条入口不完整，交付门槛都应保持未闭合。")
    builder.heading("3.2  默认目录", 2)
    builder.table(
        ["目录/对象", "用途", "管理原则"],
        [
            ["runtime_data", "设备、采集、项目与运行数据", "定期备份，不在采集中移动"],
            ["raw/high-frequency", "高频原始或标准化输入", "只读留存，禁止用表格软件覆盖原文件"],
            ["results", "处理窗口、谱结果与分析产物", "按项目/批次分层"],
            ["exports", "报告、CSV、证据包与交付清单", "发布前在独立目录复核"],
            ["logs", "运行与故障诊断", "与问题发生时间、设备 UID 一起归档"],
        ],
        [50, 65, 61],
        "工作区数据对象",
    )
    builder.callout("路径规则", "正式项目优先使用短、稳定、可备份的本地路径。网络盘可用于归档，但高频实时写入前应验证断网、锁文件和长路径行为。", "info")


def add_chapter_4(builder: ManualBuilder) -> None:
    builder.heading("4  项目、站点与元数据", 1)
    builder.paragraph("元数据不是报告装饰，而是算法输入。坐标旋转、谱修正、密度修正、足迹和单位换算都可能直接依赖它。")
    builder.figure(screenshot("04-page-project_site.png"), "项目与站点工作区。截图为演示项目。", alt="项目与站点页面显示项目概览、元数据目录和质量状态")
    builder.heading("4.1  必填信息", 2)
    builder.table(
        ["类别", "关键字段", "为什么重要"],
        [
            ["项目", "项目代码、名称、负责人、时区", "绑定批次、时间解释和交付归属"],
            ["站点", "站点代码、经纬度、海拔、地形", "足迹、气压解释和跨站比较"],
            ["下垫面", "冠层高度、粗糙度、零平面位移、地类", "足迹模型与代表性判断"],
            ["超声风速仪", "型号、序列号、轴向、安装高度、北偏角", "旋转、风向与横风修正"],
            ["气体分析仪", "类型、序列号、开/闭路、校准信息", "密度修正、诊断与谱响应"],
            ["几何布局", "水平/垂直分离、路径长度、方向", "时滞和高频衰减"],
            ["闭路管路", "长度、内径、流量、材质、加热", "管路输送时滞与 H2O 高频衰减"],
            ["时间系统", "采样率、时区、时钟源、时间戳归属", "同步、窗口切分与汇总"],
            ["输出策略", "单位、缺测码、命名和 schema", "下游兼容和重复处理"],
        ],
        [30, 73, 73],
        "科学处理所需元数据",
    )
    builder.heading("4.2  仪器布局与坐标", 2)
    builder.paragraph("记录超声风速仪坐标轴、气体分析仪相对位置和安装朝向。水平分离会导致相位差和高频损失；垂直分离会使两个传感器采样不同高度的湍涡；错误的北偏角会污染风向扇区、平面拟合和足迹叠加。")
    builder.figure(screenshot("10-project-instrument_layout.png"), "仪器布局子页。填写值必须来自现场测量或安装记录。", alt="仪器布局页面显示传感器位置和几何字段")
    builder.heading("4.3  采样链与时间", 2)
    builder.bullets([
        "开路系统重点核对光程、安装倾角、加热影响、诊断字和雨滴/结露。",
        "闭路系统重点核对管长、内径、流量、泵状态、过滤器和管路温度。",
        "各通道必须使用同一时间基准；若数据记录器已做延迟补偿，应避免二次补偿。",
        "时间戳究竟代表窗口开始、中心还是结束，必须在输出模板中固定。",
        "夏令时地区应优先存储 UTC，并在展示层转换本地时间。",
    ])
    builder.callout("现场变更", "更换仪器、改变高度、移动传感器、改管路或改变采样率后，应关闭旧方法适用期并建立新的配置版本。", "danger")


def add_chapter_5(builder: ManualBuilder) -> None:
    builder.heading("5  设备接入与实时采集", 1)
    builder.heading("5.1  设备中心", 2)
    builder.paragraph("设备中心负责发现接口、连接设备、读取参数、发送受控命令并记录应答。连接成功只说明通信链建立，不代表测量值或配置一定正确。")
    builder.steps([
        "选择设备类型、接口和地址，确认不会占用其他软件正在使用的串口或网络端口。",
        "连接后先读取设备标识、模式、采样/平均参数与诊断状态。",
        "将读回结果与站点配置、校准证书和现场记录逐项比对。",
        "需要写入时，一次只改一个参数；保存命令、应答和再次读回结果。",
        "启动采集后观察状态稳定性，再进入实时采集页。",
    ])
    builder.table(
        ["状态", "含义", "动作"],
        [
            ["在线", "通信应答正常", "继续检查数据与诊断"],
            ["采集中", "高频帧持续进入缓存和存储", "检查采样率与时间戳"],
            ["警告", "信号或参数异常但数据仍可能存在", "定位原因并标记影响时段"],
            ["故障", "通信、格式、设备诊断或存储失败", "停止交付，保留日志并处理"],
        ],
        [26, 76, 74],
        "设备状态解释",
    )

    builder.heading("5.2  实时采集", 2)
    builder.figure(screenshot("03-page-realtime.png"), "实时采集工作区。曲线和数值为演示数据。", alt="实时采集页面显示气体浓度、温压和高频流状态")
    builder.paragraph("实时曲线用于发现掉线、饱和、阶跃、漂移和不同步，不应直接当作通量。通量必须在同步、去趋势、旋转和必要修正后由协方差得到。")
    builder.table(
        ["检查项", "正常表现", "异常表现", "建议"],
        [
            ["帧率", "接近项目采样率，短时波动小", "持续低于目标或周期性归零", "查接口带宽、设备输出和磁盘"],
            ["时间戳", "单调递增，间隔接近 1/fs", "回退、重复、跳秒", "核对时钟源和解析器"],
            ["CO2/H2O", "有湍流波动且无长期饱和", "平直、断崖、超范围", "查信号强度、光路和单位"],
            ["风速分量", "u/v/w 均有合理变化", "w 恒定、符号异常、幅值失真", "查轴向、超声诊断和结冰"],
            ["温度/压力", "变化连续且量级正确", "单位错 10/100/1000 倍", "核对字段单位和传感器来源"],
            ["缓存/写盘", "持续增长且文件可读", "冻结、反复重建、磁盘满", "检查权限、空间和安全软件"],
        ],
        [27, 55, 55, 39],
        "实时质量检查",
    )
    builder.callout("异常时段", "不要删除原始样本来让曲线更平滑。记录异常起止时间和原因，由处理阶段的筛选与 QC 显式标记。", "warning")


def add_chapter_6(builder: ManualBuilder, diagrams: dict[str, Path]) -> None:
    builder.heading("6  EC 处理：方法、参数与科学含义", 1)
    builder.paragraph("EC 处理页按物理顺序组织窗口采样、清洗筛选、时滞、旋转、去趋势、协方差、密度修正、稳态、湍流、不确定度、足迹、谱修正和输出。初次配置应从站点元数据出发，而不是为了得到更好看的结果反复调参。")
    builder.figure(screenshot("05-page-ec_processing.png"), "EC 处理工作台总览。截图参数与数值为演示用途。", alt="EC 处理工作台包含步骤目录、配置区和执行状态")

    builder.heading("6.1  数据源、时间范围与预检", 2)
    builder.paragraph("数据源可选择当前项目高频目录、最近归档批次或回放文件夹。时间范围用于限定处理输入，不改变文件内容。预检应确认字段映射、单位、时间范围、有效样本、采样率、项目元数据与输出目录。")
    builder.bullets([
        "采样率优先来自可靠元数据；仅在缺失时使用时间戳推断，并检查推断稳定性。",
        "CO2、H2O、u、v、w、温度、压力应有明确单位；禁止只凭数值量级猜单位后直接交付。",
        "若样本不足以构成最小窗口，软件会不生成窗口或标记数据不足。",
        "回放文件夹适合复现问题，必须记录原始来源与复制校验信息。",
    ])

    builder.heading("6.2  窗口采样", 2)
    builder.formula("理论样本数", "N = f_s × T", "f_s 为采样率，T 为窗口秒数；有效样本数还需扣除缺测和无效帧。")
    builder.table(
        ["参数", "常见起点", "改变后的影响"],
        [
            ["采样率", "10 Hz 或仪器实际值", "过低会丢失高频贡献；填写错误会影响时滞与谱轴"],
            ["窗口长度", "30 min", "变短增加随机误差，变长增加非稳态与低频漂移风险"],
            ["缺测策略", "辅助变量线性插补/整窗保留缺测", "核心高频标量不宜用长段插值替代真实脉动"],
            ["最小有效率", "按项目 SOP", "阈值越严，保留窗口越少但完整性更高"],
        ],
        [36, 45, 95],
        "窗口采样参数",
    )

    builder.heading("6.3  数据清洗与统计筛选", 2)
    builder.paragraph("清洗用于识别数据质量问题，不应制造新的湍流结构。尖峰检测、幅值范围、绝对限值、诊断字、恒值段和缺测必须以通道为单位记录。辅助温压可在非常短的缺口内插值，核心 w、CO2、H2O 的长缺口应保留为缺测或剔除窗口。")
    builder.table(
        ["问题", "检测依据", "处理原则"],
        [
            ["尖峰", "邻域偏差、MAD/标准差、变化率", "标记并使用有界替换或缺测；保留计数"],
            ["超范围", "物理范围与仪器量程", "优先判断单位和解析错误，不直接截断"],
            ["恒值", "长时间零方差", "检查设备冻结、输出模式或解析通道"],
            ["诊断失败", "设备状态字、信号强度、光路状态", "按严重度标记窗口并保留原诊断"],
            ["时间缺口", "时间戳间隔和序列连续性", "计算缺测率，避免跨缺口形成伪协方差"],
        ],
        [29, 64, 83],
        "清洗与筛选原则",
    )

    builder.heading("6.4  时滞补偿", 2)
    builder.paragraph("风速与标量测量可能因空间分离、管路输送、仪器响应和记录链产生延迟。时滞补偿的目标是让相同气团的 w 与标量对齐。")
    builder.figure(diagrams["lag"], "在合理搜索窗内用协方差峰值估计时滞。", alt="协方差随时滞变化并在 0.55 秒达到峰值")
    builder.figure(screenshot("19-ec-lag.png"), "EC 处理的时滞设置子页。截图为演示配置。", alt="时滞设置包含搜索策略、搜索窗和期望时滞")
    builder.table(
        ["策略", "适用情况", "风险与检查"],
        [
            ["协方差最大", "信噪比足够、峰形清晰", "弱通量可能选中噪声峰；检查峰值置信度"],
            ["协方差最大带默认", "需要自动搜索并限制边界失败", "峰在边界时回退期望时滞；应复核搜索窗"],
            ["固定滞后", "几何和输送稳定、低通量或痕量气体", "设备/流量变化后固定值可能失效"],
            ["无滞后", "仅用于诊断或已在上游严格补偿", "通常不建议用于最终通量"],
        ],
        [42, 66, 68],
        "时滞策略",
    )
    builder.callout("边界峰值", "若最优时滞持续落在搜索窗最小值或最大值，先检查符号、字段同步、搜索窗、流量和管路，而不是无限扩大搜索窗。", "danger")

    builder.heading("6.5  坐标旋转", 2)
    builder.paragraph("坐标旋转把超声风速仪坐标对齐到平均流线，减少安装倾斜对垂直通量的污染。旋转前必须确认轴向、北偏角与风速符号。")
    builder.table(
        ["方法", "含义", "使用建议"],
        [
            ["双旋转", "先消除平均横风，再消除平均垂直风", "地形较平坦、窗口级处理的常用起点"],
            ["三重旋转", "在双旋转后增加第三次旋转约束", "不适用于横向应力不应为零的场景；需谨慎"],
            ["平面拟合", "用较长时期的平均风向量拟合流线平面", "复杂地形或长期站点优先；需稳定安装和足够扇区样本"],
            ["不旋转", "保留仪器坐标", "仅用于诊断、校准或已有外部旋转的输入"],
        ],
        [35, 72, 69],
        "坐标旋转方法",
    )
    builder.figure(screenshot("20-ec-rotation.png"), "坐标旋转设置子页。截图为演示配置。", alt="坐标旋转页面展示旋转方法和诊断")
    builder.callout("当前实现边界", "当前版本的单窗口平面拟合路径在无法获得长期扇区回归矩阵时会回退到双旋转。报告中的 applied_rotation_impl 与 reason 必须一起检查，不能只看下拉框选择。", "warning")

    builder.heading("6.6  横风与声温修正", 2)
    builder.paragraph("部分超声风速仪的声温或风分量会受横风、探头几何和特定型号特性影响。软件提供受控修正入口，需要匹配制造商、型号和系数来源。没有可靠型号与标定依据时应关闭，而不是套用相近型号。")
    builder.callout("系数来源", "任何自定义横风系数都应附带证书、文献或现场标定编号，并在方法来源中记录。", "warning")

    builder.heading("6.7  去趋势", 2)
    builder.table(
        ["方法", "计算思想", "影响"],
        [
            ["块均值", "减去整个窗口均值", "保留较多低频贡献，最直接满足窗口内脉动均值为零"],
            ["线性去趋势", "减去窗口内线性趋势", "抑制缓慢漂移，但可能移除真实低频通量"],
            ["滑动均值", "用有限窗口动态均值形成高通", "时间常数越小，低频损失越强"],
            ["指数滑动均值", "指数权重的动态均值", "响应连续但具有方法依赖的低频传递特性"],
        ],
        [36, 72, 68],
        "去趋势选项",
    )
    builder.paragraph("去趋势方法必须与平均窗口、谱修正和对比研究保持一致。为追求更高通量而切换去趋势会引入方法偏差。")

    builder.heading("6.8  协方差与通量", 2)
    builder.formula("离散协方差", "cov(w,c) = Σ[(w_i-w̄)(c_i-c̄)] / N", "软件在有效、同步、同一窗口内的样本上计算；具体分母与权重由方法决定。")
    builder.table(
        ["协方差方法", "说明", "使用边界"],
        [
            ["标准协方差", "对所有有效样本等权计算", "默认科研处理起点"],
            ["稳健协方差", "降低极端点对估计的支配", "用于异常敏感诊断；需说明与标准结果差异"],
            ["窗口内加权", "按窗口内规则赋权", "权重必须有物理和统计依据"],
        ],
        [42, 68, 66],
        "协方差模式",
    )

    builder.heading("6.9  密度修正与混合比", 2)
    builder.paragraph("气体分析仪可能输出摩尔密度、摩尔分数或混合比。热和水汽引起的空气密度涨落会使原始浓度协方差偏离真实标量交换。经典密度修正把水汽与温度项加入开路密度通量；闭路系统在有可靠池温池压时通常优先转换为混合比。")
    builder.formula("概念表达", "F_corrected = F_raw + F_water-density + F_heat-density", "实际项取决于分析仪类型、变量表达和可用温压元数据；不得在未知输入类型上盲目套用。")
    builder.table(
        ["选项", "适用条件", "必须核对"],
        [
            ["WPL", "以密度形式测量且需要补偿热/水汽密度效应", "温度、压力、水汽单位、分析仪类型"],
            ["混合比优先", "闭路池温池压可信，可转换混合比", "池体温压来源、同步与诊断"],
            ["不修正", "输入已在上游完成等效处理或仅做诊断", "避免遗漏或重复修正"],
        ],
        [38, 74, 64],
        "密度处理策略",
    )
    builder.callout("修正幅度", "密度修正项若相对原始协方差异常巨大，优先检查单位、温压来源、H2O 通道、时滞和符号。大修正不等于大真实通量。", "danger")

    builder.heading("6.10  稳态与湍流", 2)
    builder.paragraph("稳态检验比较完整窗口协方差与多个子窗口协方差；当前实现使用六个子窗口。湍流检验综合 u*、风速方差与平均风速等信息。数据不足时应得到 insufficient_data，而不是伪造一个正常分数。")
    builder.formula("摩擦速度", "u* = [(overline(u′w′))² + (overline(v′w′))²]^(1/4)", "u* 描述动量交换强度；夜间低 u* 常对应间歇湍流和通量低估风险。")
    builder.figure(screenshot("25-ec-steadiness.png"), "稳态检验设置与结果区域。截图为演示数据。", alt="稳态检验页面显示子窗口对比和质量状态")
    builder.bullets([
        "稳态差异小只说明窗口内一致性较好，不证明仪器校准或足迹正确。",
        "u* 阈值具有站点、季节和冠层依赖性；经验默认值只能作为起点。",
        "夜间弱湍流筛选应与生态分析目标、存储项和缺测填补策略共同制定。",
        "低风速下旋转和风向本身也更不稳定，应联动解释。",
    ])

    builder.heading("6.11  随机不确定度", 2)
    builder.paragraph("软件提供 Mann-Lenschow、Finkelstein-Sims 与组合经验方法。前两者从湍流时间相关或协方差统计估计随机误差；组合经验方法把时滞置信、稳态、湍流、连续性和密度修正幅度形成审计性风险带。不同方法不是可互换的同一数值。")
    builder.figure(screenshot("27-ec-uncertainty.png"), "不确定度方法设置。截图为演示参数。", alt="不确定度页面显示方法、积分时间尺度和置信水平")
    builder.table(
        ["方法", "需要的信息", "解释"],
        [
            ["Mann-Lenschow", "样本数、平均期、积分时间尺度", "关注有限采样导致的随机误差"],
            ["Finkelstein-Sims", "w 与标量协方差序列、采样率", "从协方差函数估计采样误差"],
            ["组合经验", "QC 组件与修正幅度", "用于运行风险概览，不替代严格统计不确定度"],
        ],
        [42, 66, 68],
        "不确定度方法",
    )

    builder.heading("6.12  通量足迹", 2)
    builder.figure(diagrams["footprint"], "足迹峰值与累计贡献区域的概念示意。", alt="观测塔下风方向的多个足迹贡献椭圆")
    builder.table(
        ["方法", "特点", "适用与限制"],
        [
            ["Kljun", "缩放参数化，可生成 2D 足迹", "需要测量高度、稳定度、u* 等；复杂地形慎用"],
            ["Kormann-Meixner", "基于平流扩散与幂律剖面", "对输入剖面与稳定度敏感"],
            ["Hsieh", "基于稳定度分区的解析近似", "适合快速估计；应说明模型简化"],
        ],
        [42, 60, 74],
        "足迹模型",
    )
    builder.paragraph("启用网格输出时，可导出二维网格、等值线、GeoJSON、GeoTIFF 及地类叠加证据。地类栅格坐标系、分辨率与覆盖范围必须验证；没有重叠时不得强行解释来源比例。")

    builder.heading("6.13  谱修正与方法对比", 2)
    builder.paragraph("谱修正用于估计有限频率响应造成的通量损失。当前版本支持 Massman、Horst、Ibrom 和 Fratini 方法族；Fratini 路径可在有匹配协谱时使用实测协谱证据。路径平均、传感器分离、响应时间、测量高度和稳定度等参数必须来自项目元数据。")
    builder.table(
        ["方法族", "主要思路", "使用建议"],
        [
            ["Massman", "组合系统传递函数估计频率损失", "适合建立组件化衰减预算"],
            ["Horst", "一阶响应标量传感器的简化衰减表达", "参数少，适合敏感性检查"],
            ["Ibrom", "关注闭路水汽等低通与时变衰减", "湿度、管路和流量信息要可靠"],
            ["Fratini", "用实测/参考协谱与传递信息估计修正", "必须检查协谱匹配来源和质量"],
        ],
        [36, 76, 64],
        "谱修正方法族",
    )
    builder.paragraph("方法对比功能用于暴露模型敏感性，不应从多个方法中挑选“结果最大”的一个。差异超过项目阈值时，应回到几何、稳定度、协谱和参数来源进行解释。")

    builder.heading("6.14  输出模式", 2)
    builder.paragraph("only_available 只输出本批次实际可计算字段；standard_schema 保持标准列集合，并以明确缺测表示不可用字段。前者紧凑，后者利于长期数据库和跨批次拼接。无论哪种模式，都应保存单位、缺测码、方法快照和字段来源。")
    builder.callout("可用不等于可信", "字段存在只说明算法产生了值。是否适合科研使用，还要结合 QC、诊断、足迹、不确定度和方法限制。", "warning")


def add_chapter_7(builder: ManualBuilder, diagrams: dict[str, Path]) -> None:
    builder.heading("7  谱修正与 QC 工作区", 1)
    builder.paragraph("谱工作区把窗口级频域证据放在一起，用于判断通量主要来自哪些时间尺度、传感器是否漏掉高频或低频贡献、修正因子是否有物理依据。")
    builder.figure(screenshot("06-page-spectral_qc.png"), "谱修正与 QC 工作区总览。图中结果为演示数据。", alt="谱质控页面包含窗口列表、图形预览和质量诊断")
    builder.heading("7.1  功率谱与协谱", 2)
    builder.paragraph("功率谱描述单个变量在不同频率上的方差分布；协谱描述 w 与标量在各频率对通量协方差的贡献。通量修正应主要依据协谱，因为两个变量各自有能量并不保证它们共同贡献净通量。")
    builder.formula("协方差与协谱", "overline(w′c′) = ∫ Co_wc(f) df", "理想情况下，协谱对频率积分应与时域协方差一致；有限采样和离散算法会带来容差。")
    builder.figure(screenshot("32-qc-cross_spectrum.png"), "协谱页面。检查贡献符号、峰值频段和高频尾部。演示数据。", alt="协谱图显示垂直风速与标量的频率贡献")
    builder.table(
        ["现象", "可能原因", "复核动作"],
        [
            ["高频快速衰减", "管路、响应时间、路径平均、空间分离", "核对几何、流量和传递函数组件"],
            ["低频端不收敛", "窗口过短、非稳态、漂移、风向改变", "看 Ogive、稳态子窗和原始时序"],
            ["协谱多峰或符号交替", "信噪比低、时滞错误、间歇事件", "检查时滞曲线、诊断字和异常时段"],
            ["功率谱正常但协谱接近零", "标量波动与 w 不相关或通量很小", "不要仅凭浓度波动判定有通量"],
            ["Nyquist 附近抬升", "电子噪声、混叠或解析异常", "核对采样率、数字滤波和原始帧"],
        ],
        [42, 68, 66],
        "常见谱形诊断",
    )

    builder.heading("7.2  Ogive 累积曲线", 2)
    builder.paragraph("Ogive 是从高频或低频方向累积协谱得到的曲线。理想情况下，曲线在包含主要湍涡后趋于平台；持续斜坡、反向或末端大幅跳动提示窗口、趋势、时滞或非稳态问题。")
    builder.figure(screenshot("33-qc-ogive.png"), "Ogive 页面。平台收敛应与稳态和窗口长度共同解释。演示数据。", alt="Ogive 累积协谱曲线")
    builder.bullets([
        "平台出现过早可能意味着高通处理过强或高频贡献受限。",
        "低频末端仍显著变化说明平均窗口可能未覆盖主要尺度，或存在趋势和天气转变。",
        "正负反复穿越不一定是错误，但需要查看协谱频段与原始事件。",
        "不要用视觉平滑替代数值积分与一致性检查。",
    ])

    builder.heading("7.3  传递函数与修正因子", 2)
    builder.figure(diagrams["spectral"], "理想与观测协谱之间的高频衰减示意。", alt="理想协谱和观测协谱在高频端分离")
    builder.figure(screenshot("34-qc-transfer_function.png"), "传递函数页面。各组件应能追溯到仪器和布局参数。演示数据。", alt="传递函数页面展示总响应和组件贡献")
    builder.paragraph("总传递函数通常由多个组件相乘。某一组件参数错误会在整个频段传播，因此应查看路径平均、空间分离、传感器响应、管路和相位等分量，而不只看最终修正因子。")
    builder.table(
        ["组件", "主要输入", "典型影响"],
        [
            ["路径平均", "超声路径、光程、平均风速", "衰减小尺度高频波动"],
            ["空间分离", "传感器水平/垂直间距、风速", "造成相位差与高频损失"],
            ["传感器响应", "时间常数或响应带宽", "一阶或更复杂低通"],
            ["闭路管路", "长度、内径、流量、湿度、温度", "时滞与标量相关的高频衰减"],
            ["数字处理", "采样、平均、滤波、保持", "改变 Nyquist 附近响应"],
            ["去趋势/窗口", "时间常数与平均期", "影响低频端贡献"],
        ],
        [36, 68, 72],
        "频率响应组件",
    )
    builder.callout("过度修正", "修正因子很大、随窗口剧烈跳变或与风速/湿度没有合理关系时，应判为风险证据。先修正元数据和时滞，再决定是否接受该窗口。", "danger")

    builder.heading("7.4  窗口审阅顺序", 2)
    builder.steps([
        "从 QC 总览筛出 C 级、修正因子异常、时滞置信低和数据不足窗口。",
        "检查窗口详情与原始质量：样本数、缺测、诊断字、风向与事件。",
        "检查时滞/相位：峰值是否在合理窗内，CO2 与 H2O 是否呈可解释差异。",
        "检查功率谱与协谱：高频尾部、低频趋势、噪声与符号。",
        "检查 Ogive：主要贡献是否收敛，末端是否受非稳态支配。",
        "检查传递函数和修正因子：组件与元数据是否一致。",
        "记录结论：接受、带限制接受、重新处理或排除，并写明原因。",
    ])


def add_chapter_8(builder: ManualBuilder) -> None:
    builder.heading("8  QC、结果与科学解释", 1)
    builder.paragraph("结果解释应从“数据是否完整”开始，经过“方法是否适用”和“物理证据是否一致”，最后才讨论通量大小。单一总分不能覆盖所有风险。")
    builder.heading("8.1  QC 矩阵", 2)
    builder.table(
        ["质量维度", "关注问题", "异常含义"],
        [
            ["信号有效性", "设备诊断、量程、信号强度", "输入本身可能失真"],
            ["连续性", "缺测、时间缺口、有效率", "协方差采样不完整"],
            ["时滞置信", "峰值清晰度与边界", "w 与标量可能未正确对齐"],
            ["旋转", "实施方法、平均 w 与回退", "坐标可能未对齐流线"],
            ["稳态", "全窗与子窗协方差差异", "窗口内过程可能变化"],
            ["湍流", "u*、方差、平均风速", "交换可能不足或间歇"],
            ["密度修正", "修正幅度与输入来源", "单位/温压/H2O 或方法可能异常"],
            ["频率响应", "协谱、Ogive、修正因子", "观测可能漏失尺度贡献"],
            ["足迹", "来源距离、地类和适用性", "结果代表区域可能不符合目标"],
        ],
        [37, 75, 64],
        "窗口质量维度",
    )
    builder.heading("8.2  A/B/C 与 0/1/2", 2)
    builder.paragraph("软件可将稳态与湍流的分级结果映射为 A/B/C，其中 A 对应高质量、B 表示可用于部分预算或需复核、C 表示不建议直接用于最终数据集。需要与外部网络兼容时，可按项目方法映射为 0/1/2。")
    builder.table(
        ["内部等级", "常见映射", "建议用途", "必要动作"],
        [
            ["A", "0", "主要科研与预算分析候选", "仍检查诊断、足迹和异常事件"],
            ["B", "1", "可在限定用途使用", "保留限制说明，必要时做敏感性分析"],
            ["C", "2", "通常排除最终聚合", "定位原因；可保留作诊断证据"],
        ],
        [28, 30, 65, 53],
        "质量等级的使用",
    )
    builder.callout("等级不是结论", "A 级不保证足迹目标正确或能量闭合；C 级也不应从原始档案中删除。质量等级是可重复决策的输入。", "warning")

    builder.heading("8.3  通量符号与量级", 2)
    builder.bullets([
        "先确认垂直轴向上为正、风速分量映射正确，再解释正负号。",
        "CO2 白天下垫面吸收常表现为负值，夜间呼吸常为正值，但具体生态系统与事件可能不同。",
        "潜热通量与 H2O 通量需要一致的摩尔/质量单位和汽化潜热换算。",
        "显热通量依赖声温到空气温度的处理和湿度影响，必须检查温度来源。",
        "量级异常时先查单位、时间基准和修正项，不以历史期望值强行裁剪。",
    ])
    builder.heading("8.4  不确定度、足迹与用途", 2)
    builder.paragraph("随机不确定度描述有限湍流采样的不稳定性；系统误差还来自校准、代表性、频率响应和模型假设。足迹则描述空间来源。一个数值可有较小随机误差，但来源区域并非目标地类；也可足迹良好但弱湍流导致通量低估。报告应把三者分开。")
    builder.table(
        ["结果组合", "解释", "建议"],
        [
            ["低随机不确定度 + A + 目标足迹", "综合证据较强", "进入最终候选并做聚合 QA"],
            ["低不确定度 + 非目标足迹", "数值稳定但代表性不符", "按风向/地类筛选或单列"],
            ["高不确定度 + A", "质量测试通过但随机采样不足", "扩大时间聚合并保留误差"],
            ["C + 目标足迹", "来源正确但统计条件不足", "通常不用于最终预算"],
        ],
        [58, 66, 52],
        "多证据联合解释",
    )


def add_chapter_9(builder: ManualBuilder) -> None:
    builder.heading("9  报告、证据与交付", 1)
    builder.paragraph("报告中心把运行结果、质量诊断、方法来源和交付文件组织成可审阅包。导出不是简单复制 CSV，而是一次科学与工程闭环。")
    builder.figure(screenshot("07-page-report_center.png"), "报告中心总览。截图为演示数据。", alt="报告中心包含报告目录、预览、交付门槛和导出操作")
    builder.heading("9.1  报告目录", 2)
    builder.table(
        ["报告页", "主要内容", "审核重点"],
        [
            ["运行摘要", "项目、批次、时间范围、状态和 KPI", "输入批次与版本是否正确"],
            ["设备状态", "在线、采集、告警与诊断", "异常时段是否影响结果"],
            ["采集质量", "帧率、缺测、连续性、通道可用性", "样本完整性"],
            ["EC 结果", "窗口通量、均值、QC 与修正", "单位、符号和方法"],
            ["谱修正与 QC", "谱证据、修正因子与窗口详情", "修正是否有物理依据"],
            ["异常事件", "错误、警告和处理记录", "未闭合事件是否仍存在"],
            ["站点与方法", "项目元数据和处理配置", "方法适用期与参数来源"],
            ["证据包", "文件、清单、校验与导出状态", "交付完整性"],
            ["方法来源", "方法选择、实现路径和限制", "实际执行方法而非仅配置值"],
            ["方法对比", "方法族差异和阈值", "差异原因与接受标准"],
        ],
        [39, 75, 62],
        "报告中心内容",
    )
    builder.heading("9.2  生成报告", 2)
    builder.steps([
        "选择正确项目、批次和视图模式。",
        "点击生成报告，确认预览来自真实运行结果而非空状态。",
        "逐页查看数据、证据与结论，重点复核 C 级窗口和异常修正。",
        "检查报告版本、来源批次、时间范围、方法和限制。",
        "解决交付门槛中的未闭合项后重新生成，避免报告与最终结果不一致。",
    ])
    builder.heading("9.3  导出报告与证据包", 2)
    builder.paragraph("根据可用结果，导出可能包含完整输出 CSV、窗口与 QC 明细、谱评估、足迹网格/GIS 文件、方法对比、运行清单、网络或 schema 校验以及正式报告。某类输入或方法未启用时，对应产物可以明确标记不可用，而不应伪造空结果。")
    builder.figure(screenshot("45-report-evidence_pack.png"), "证据包页：核对清单、验证和归档状态。演示数据。", alt="证据包页面展示交付文件和验证状态")
    builder.figure(screenshot("46-report-method_provenance.png"), "方法来源页：检查实际实现、参数来源和限制。演示数据。", alt="方法来源页面展示旋转、足迹、不确定度和谱方法的溯源")
    builder.table(
        ["交付门槛", "通过条件", "失败时"],
        [
            ["报告", "预览包含真实批次和关键章节", "重新生成或先运行处理"],
            ["导出", "目标文件落盘且可打开", "检查路径、权限和磁盘空间"],
            ["清单", "文件名、大小、角色与校验信息完整", "重新构建证据包"],
            ["网络/schema", "字段、单位和缺失策略满足约定", "修正模板或记录例外"],
            ["方法", "实际方法、参数、来源和限制可追溯", "补元数据或重跑"],
            ["异常", "未闭合问题已清零或有正式豁免", "禁止直接发布"],
        ],
        [36, 82, 58],
        "交付前检查",
    )
    builder.callout("独立复核", "正式交付后应在新的空目录中解包/打开，验证报告、CSV、GIS 和清单相互一致。不要只在生成机器上检查。", "success")


def add_chapter_10(builder: ManualBuilder) -> None:
    builder.heading("10  标准作业程序（SOP）", 1)
    builder.heading("10.1  新站点启用", 2)
    builder.steps([
        "建立项目和站点代码，固定时区、坐标和文件命名。",
        "录入仪器型号、序列号、校准、固件、轴向和安装高度。",
        "测量传感器几何与闭路管路参数，拍照并归档安装证据。",
        "核对设备时间、数据记录器时间和计算机时间。",
        "采集零风/静态或设备自检数据，确认风速偏置和气体诊断。",
        "运行短时预检，验证字段、单位、采样率和数据落盘。",
        "采集足够长期数据后评估平面拟合、时滞分布和 u* 阈值。",
        "冻结首个方法版本并生成基线报告与证据包。",
    ])
    builder.heading("10.2  每日运行", 2)
    builder.table(
        ["时间", "动作", "记录"],
        [
            ["开始前", "设备在线、时间、磁盘、项目与天气", "值班记录"],
            ["采集中", "帧率、缓存、诊断字、曲线与告警", "异常起止时间"],
            ["处理前", "数据完整性和站点变更", "预检结果"],
            ["处理后", "QC、时滞、谱、修正、足迹", "窗口审阅记录"],
            ["结束时", "报告、证据包、备份和交付门槛", "批次清单"],
        ],
        [27, 91, 58],
        "每日运行检查",
    )
    builder.heading("10.3  历史数据重处理", 2)
    builder.steps([
        "复制或只读挂载历史原始数据，计算文件清单，不改变原始档案。",
        "建立新的处理批次和方法版本，记录重处理原因。",
        "运行预检并抽查原始时序、单位和时区。",
        "使用固定配置处理全部范围；不要在看到结果后逐窗口手调参数。",
        "对新旧结果做窗口配对比较，按旋转、时滞、密度、谱和 QC 分解差异。",
        "形成变更说明、影响范围和回滚方案，再替换下游数据集。",
    ])
    builder.heading("10.4  正式交付", 2)
    builder.steps([
        "冻结输入清单、项目元数据和处理配置。",
        "运行完整处理并保存软件版本与实际方法实现。",
        "完成异常窗口、谱证据、足迹和不确定度审阅。",
        "生成报告、标准输出和证据包。",
        "检查交付门槛、缺测码、单位、时区、字段 schema 与文件校验。",
        "在独立目录复核打开与可读性，记录审核人和发布日期。",
        "将原始数据、配置、结果、报告与日志分层归档。",
    ])
    builder.callout("配置冻结", "同一交付批次中，任何影响计算的参数变化都应触发重新处理和报告重生。不要把不同配置的窗口拼成一个未标注数据集。", "danger")


def add_chapter_11(builder: ManualBuilder) -> None:
    builder.heading("11  故障排查", 1)
    builder.paragraph("排查顺序固定为：确认项目与时间 -> 查看底部日志 -> 检查设备/数据 -> 检查方法预检 -> 检查窗口证据 -> 最后检查导出。一次只改变一个变量，并保留前后对比。")
    builder.table(
        ["现象", "常见原因", "诊断步骤", "处理"],
        [
            ["软件启动闪烁或界面反复刷新", "显示缩放、驱动、状态更新过密", "记录缩放比例、显示器、日志和复现步骤", "使用发布版默认缩放；更新驱动；提交日志与截图"],
            ["设备在线但无数据", "输出模式未启用、端口被占用、解析不匹配", "查设备应答、原始帧、端口和模式", "读回配置，释放端口，选择正确协议"],
            ["帧率不足", "接口带宽、设备平均、CPU/磁盘负载", "比较设备输出率、接收率和写盘率", "减少非必要日志，修正采样设置，换本地磁盘"],
            ["时间戳跳变", "时钟同步、时区、设备重启、解析", "定位首个跳变和对应日志", "修正时间源；分割受影响批次"],
            ["没有生成窗口", "数据太短、采样率错、有效率不足", "查样本数、fs、窗口长度和缺测", "修正元数据或选择足够时间范围"],
            ["时滞总在边界", "搜索窗错、通道不同步、弱通量、管路变化", "看协方差曲线和期望时滞", "修正搜索窗/流量；必要时固定滞后"],
            ["旋转后平均 w 仍大", "轴向错、低风速、复杂地形、回退", "查 alpha/beta、actual impl、风向与风速", "修正轴向；长期平面拟合；标记低风窗口"],
            ["密度修正异常大", "单位、温压、水汽、重复修正", "分解原始项、水汽项、热项", "修正输入表达和元数据后重跑"],
            ["修正因子异常大", "几何/响应参数错、协谱不匹配、噪声", "查传递组件、匹配来源和谱形", "修正参数；拒绝无依据修正"],
            ["Ogive 不收敛", "非稳态、窗口短、趋势或低频事件", "查原始时序、稳态和风向", "分割事件、调整窗口或排除"],
            ["足迹无有效输出", "高度/稳定度/u* 缺失或栅格不重叠", "查模型输入和 GIS 验证", "补元数据、修正 CRS 或标记不可用"],
            ["报告为空", "没有真实处理批次或筛选错项目", "核对项目、批次、时间范围", "先运行处理并刷新报告中心"],
            ["导出失败", "权限、长路径、磁盘、文件占用", "看错误日志和目标目录", "换短本地路径，关闭占用，释放空间"],
        ],
        [38, 47, 53, 38],
        "常见故障与处置",
    )
    builder.heading("11.1  提交问题时的最小证据", 2)
    builder.bullets([
        "软件版本、Windows 版本、显示缩放和项目代码。",
        "问题发生时间、设备 UID、输入批次和最短复现步骤。",
        "完整日志片段，不只提供弹窗文字。",
        "相关配置快照、方法来源和预检结果。",
        "可脱敏的最小原始样本或回放包，以及期望与实际结果。",
        "界面问题附全屏截图；数据问题附窗口 ID、谱/QC 证据。",
    ])
    builder.callout("不要直接覆盖", "排障期间保留失败批次。只有成功重跑并完成差异复核后，才能标记旧结果为已替代。", "warning")


def add_chapter_12(builder: ManualBuilder) -> None:
    builder.heading("12  数据治理、版本与科学限制", 1)
    builder.heading("12.1  可重复处理的最小集合", 2)
    builder.table(
        ["对象", "至少保存", "目的"],
        [
            ["原始数据", "原文件、时间范围、大小与校验", "重现输入"],
            ["项目元数据", "站点、仪器、布局、采样链、时间", "重现物理条件"],
            ["处理配置", "所有步骤参数与默认值", "重现算法选择"],
            ["软件信息", "版本、构建、发布来源", "重现实现"],
            ["运行证据", "日志、预检、窗口结果、谱与 QC", "解释过程"],
            ["交付清单", "文件角色、大小、校验和状态", "验证完整性"],
            ["审阅记录", "接受/排除理由、审核人和日期", "重现科学决策"],
        ],
        [36, 87, 53],
        "可重复性档案",
    )
    builder.heading("12.2  备份与保留", 2)
    builder.bullets([
        "采集盘与归档盘分离；采集时不在同一目录运行同步重命名。",
        "原始数据至少保留一份离线副本和一份独立位置副本。",
        "派生结果可重算，但正式交付批次仍应与配置、版本和审阅记录一起保留。",
        "定期执行恢复演练，确认备份不是只有文件名而无法读取。",
        "包含站点坐标、设备序列号或客户信息时，按项目权限和隐私规则分发。",
    ])
    builder.heading("12.3  版本信息与更新日志", 2)
    builder.paragraph("软件版本可在标题栏、底部版本标签和关于窗口查看。每次升级前阅读发布说明，重点关注处理算法、默认参数、字段 schema、导出格式和兼容性变化。升级后先用固定回放批次做回归对比，再处理正式数据。")
    builder.table(
        ["变更类型", "升级后验证"],
        [
            ["界面/性能", "启动、页面切换、长任务、显示缩放和导出"],
            ["采集/协议", "设备读回、帧率、时间戳、断线恢复和原始存储"],
            ["算法", "固定样本上的通量、时滞、QC、谱和足迹差异"],
            ["输出", "列名、单位、缺测码、清单、下游导入和历史拼接"],
            ["依赖/打包", "离线启动、权限、安全软件和可执行文件版本"],
        ],
        [52, 124],
        "升级回归清单",
    )
    builder.heading("12.4  科学限制", 2)
    builder.bullets([
        "软件不能从不充分湍流中恢复真实夜间通量，也不能自动补偿未测量的水平平流和存储项。",
        "复杂地形、城市、林缘和强异质下垫面可能超出常用旋转与足迹模型假设。",
        "谱修正依赖仪器响应和几何元数据；未知参数会形成不可量化系统误差。",
        "质量标志和不确定度是方法依赖结果，跨软件或跨版本比较前必须对齐处理配置。",
        "没有真实数据验证时，代码级测试只能证明实现行为和不变量，不能证明某站点的科学准确度。",
        "最终科研结论仍需结合能量平衡、生态/气象背景、独立仪器和长期一致性。",
    ])
    builder.callout("发布声明", "报告应区分“软件计算完成”“质量检查通过”和“科学结论成立”三个层级。只有第三层需要研究负责人对站点与研究问题作最终判断。", "danger")


def add_appendix_a(builder: ManualBuilder) -> None:
    builder.heading("附录 A  关键参数字典", 1)
    builder.paragraph("下表给出界面参数的物理含义和设置原则。实际可选值以当前软件版本和所选设备/方法为准。")
    builder.table(
        ["参数", "含义/单位", "设置原则", "错误风险"],
        [
            ["sample_hz", "高频采样率 / Hz", "使用设备真实输出率并用时间戳验证", "时滞和频率轴错误"],
            ["window_minutes", "平均窗口 / min", "常用 30 min；按站点平稳性验证", "随机误差或非稳态增加"],
            ["missing_policy", "缺测策略", "核心高频量避免长段插值", "制造伪协方差"],
            ["min_valid_ratio", "最小有效样本比例", "按项目 SOP 与用途制定", "阈值过松保留坏窗"],
            ["lag_strategy", "时滞策略", "按信噪比、管路与稳定性选择", "错位导致通量低估/反号"],
            ["search_window_s", "时滞搜索窗 / s", "覆盖物理可达延迟且不过宽", "选中噪声峰"],
            ["expected_lag_s", "期望/默认时滞 / s", "由几何、流量和长期统计得到", "错误回退"],
            ["rotation_mode", "坐标旋转方法", "平坦站点可双旋转；复杂地形评估平面拟合", "倾斜污染垂直通量"],
            ["detrend_mode", "去趋势方法", "与研究网络、窗口和谱方法一致", "低频贡献改变"],
            ["covariance_mode", "协方差估计", "标准方法为默认；稳健/加权需说明", "方法偏差"],
            ["density_mode", "密度/混合比处理", "按分析仪类型和输入表达选择", "重复或遗漏修正"],
            ["crosswind_enabled", "横风/声温修正", "仅在型号与系数有依据时启用", "套错系数"],
            ["sonic_model", "超声型号", "按铭牌、固件和证书填写", "错误修正路径"],
            ["stationarity_rule", "稳态判定", "保持项目内一致并保存子窗证据", "跨批次不可比"],
            ["ustar_rule", "u* 阈值规则", "优先使用站点季节阈值", "夜间低估或过度剔除"],
            ["footprint_enabled", "是否计算足迹", "有足够站点和湍流输入时开启", "无效代表性结论"],
            ["footprint_method", "足迹模型", "按地形、数据和研究目的选择", "模型外推"],
            ["z_m", "有效测量高度 / m", "相对位移高度定义一致", "足迹距离严重偏差"],
            ["canopy_height_m", "冠层高度 / m", "使用对应时期观测", "零平面位移错误"],
            ["roughness_length_m", "粗糙度长度 / m", "站点/风向特定值更可靠", "足迹和稳定度偏差"],
            ["obukhov_length_m", "Monin-Obukhov 长度 / m", "优先用窗口计算值", "稳定度分区错误"],
            ["grid_x/y_bins", "足迹网格分辨率", "兼顾空间细节与文件大小", "过粗丢细节，过细假精度"],
            ["uncertainty_method", "随机不确定度方法", "按统计假设和可用数据选择", "不同方法数值不可比"],
            ["integral_timescale_s", "积分时间尺度 / s", "由自相关或长期统计估计", "随机误差偏差"],
            ["confidence_level", "置信水平", "报告中固定并明确", "误读误差带"],
            ["spectral_method", "谱修正方法族", "与系统类型和证据匹配", "不适用模型修正"],
            ["path_length_m", "测量路径 / m", "按仪器几何填写", "路径平均衰减错误"],
            ["separation_m", "传感器分离 / m", "记录方向与分量", "相位/高频损失错误"],
            ["response_time_s", "传感器响应时间 / s", "来自证书或标定", "传递函数错误"],
            ["cospectrum_source", "协谱来源", "优先使用同批次匹配协谱并验证", "修正证据错配"],
            ["method_compare", "方法族对比开关", "用于敏感性与归因，不用于挑最大值", "选择性偏差"],
            ["deviation_threshold", "方法差异阈值", "按项目容忍度预先规定", "事后调阈值"],
            ["full_output_mode", "输出列策略", "长期库用标准 schema；诊断可只输出可用", "下游列漂移"],
        ],
        [36, 48, 58, 34],
        "处理配置参数",
    )


def add_appendix_b(builder: ManualBuilder) -> None:
    builder.heading("附录 B  常用输出字段与单位", 1)
    builder.paragraph("字段名可能随输出模板变化。导入下游系统前应以本批次 schema、单位行和方法清单为准。")
    builder.table(
        ["字段/概念", "典型单位", "说明"],
        [
            ["window_start / window_end", "ISO 8601", "窗口边界与时区必须明确"],
            ["sample_count", "1", "理论或进入窗口的样本数"],
            ["valid_sample_count", "1", "通过有效性检查的样本数"],
            ["continuity_ratio", "0-1", "时间连续程度"],
            ["missing_ratio", "0-1", "缺测比例"],
            ["mean_co2", "µmol mol⁻¹", "窗口平均 CO2 摩尔分数或指定表达"],
            ["mean_h2o", "mmol mol⁻¹", "窗口平均 H2O 摩尔分数或指定表达"],
            ["mean_pressure", "kPa", "密度和热力学处理使用"],
            ["mean_temperature", "°C", "温度来源需在方法中说明"],
            ["lag_seconds", "s", "最终应用时滞"],
            ["lag_confidence", "0-1", "峰值可信程度或归一化指标"],
            ["rotation_mode", "文本", "配置旋转方法"],
            ["applied_rotation_impl", "文本", "实际执行实现；可与配置不同"],
            ["alpha_deg / beta_deg", "degree", "旋转角"],
            ["cov_w_co2", "依输入单位", "w 与 CO2 的原始协方差"],
            ["cov_w_h2o", "依输入单位", "w 与 H2O 的原始协方差"],
            ["raw_flux", "通量单位", "密度/谱修正前的通量"],
            ["primary_flux", "通量单位", "当前方法选定的主通量"],
            ["primary_flux_source", "文本", "主通量来自原始、混合比或密度修正路径"],
            ["wpl_water_vapor_term", "通量单位", "水汽密度项"],
            ["wpl_sensible_heat_term", "通量单位", "热密度项"],
            ["spectral_factor", "1", "谱修正乘数或等效因子"],
            ["ustar", "m s⁻¹", "摩擦速度"],
            ["stationarity_score", "0-100", "窗口稳态指标；无数据时为空"],
            ["turbulence_score", "0-100", "湍流充分度指标；无数据时为空"],
            ["qc_grade", "A/B/C", "综合质量等级"],
            ["qc_flags / reasons", "文本", "触发的检查项和原因"],
            ["uncertainty", "通量单位或比例", "方法相关的不确定度估计"],
            ["footprint_peak_distance", "m", "最大贡献位置"],
            ["footprint_contribution_distances", "m", "10/30/50/70/90% 等累计距离"],
            ["method_provenance", "结构化文本", "方法、参数、来源、限制与实际实现"],
        ],
        [57, 36, 83],
        "核心输出字段",
    )
    builder.callout("单位优先", "字段名相同而单位不同，比字段名不同更危险。任何外部拼接都应先验证单位、时区、缺测码和变量表达。", "warning")


def add_appendix_c(builder: ManualBuilder) -> None:
    builder.heading("附录 C  QC 决策矩阵", 1)
    builder.table(
        ["场景", "状态", "默认决策", "可接受例外"],
        [
            ["设备诊断失败或核心通道饱和", "红", "排除窗口", "无；仅作故障诊断"],
            ["时间戳回退/大缺口", "红", "分割批次并重跑", "无可靠重建依据时排除"],
            ["时滞峰在边界且置信低", "红", "修正时滞配置后重跑", "固定滞后有长期证据时可用"],
            ["密度修正项远大于原始项", "红", "核对单位/温压/H2O", "强蒸发事件也需独立证据"],
            ["谱修正因子超项目上限", "红", "检查传递函数并拒绝盲修正", "方法验证后带限制保留"],
            ["C 级稳态/湍流", "红", "通常排除最终数据集", "方法学研究可保留"],
            ["B 级且其他证据正常", "黄", "带标志进入候选", "按用途和网络规则决定"],
            ["足迹主要落在非目标地类", "黄", "按风向/地类筛选", "研究问题包含该地类时可单列"],
            ["随机不确定度高", "黄", "保留误差并谨慎聚合", "更长聚合可能降低随机误差"],
            ["A 级、谱收敛、足迹目标一致", "绿", "进入最终候选", "仍需总体 QA 与外部一致性"],
        ],
        [55, 18, 62, 41],
        "窗口接受与排除建议",
    )
    builder.heading("附录 C.1  审阅记录模板", 2)
    builder.table(
        ["字段", "填写要求"],
        [
            ["项目/站点/批次", "唯一识别信息"],
            ["窗口 ID 与时间", "含时区"],
            ["触发项", "QC、诊断、谱、足迹或异常事件"],
            ["证据", "图、字段、日志或外部记录路径"],
            ["决策", "接受 / 带限制接受 / 重处理 / 排除"],
            ["理由", "可复核的物理或统计依据"],
            ["审核人/日期", "责任与时间"],
            ["后续动作", "重跑批次、配置版本或报告编号"],
        ],
        [47, 129],
        "人工审阅最小记录",
    )


def add_appendix_d(builder: ManualBuilder) -> None:
    builder.heading("附录 D  术语表", 1)
    builder.table(
        ["术语", "说明"],
        [
            ["EC", "Eddy Covariance，涡动协方差法"],
            ["Reynolds 分解", "把瞬时量分为平均量与脉动量"],
            ["协方差", "两个脉动量共同变化的平均乘积"],
            ["功率谱", "单变量方差在频率上的分布"],
            ["协谱", "两变量协方差在频率上的分布"],
            ["Ogive", "协谱随频率累积得到的通量贡献曲线"],
            ["传递函数", "测量系统对不同频率信号的幅值/相位响应"],
            ["频率响应修正", "估计测量系统漏失频率贡献的过程"],
            ["时滞", "风速与标量对同一气团响应的时间差"],
            ["去趋势", "从时序移除均值或慢变化以定义脉动"],
            ["双旋转", "使平均横风和平均垂直风接近零的两步旋转"],
            ["三重旋转", "在双旋转基础上增加横向应力约束"],
            ["平面拟合", "用长期平均风矢量拟合倾斜流线平面"],
            ["WPL", "对热和水汽引起的空气密度涨落进行补偿"],
            ["混合比", "标量物质的量与干空气物质的量之比"],
            ["u*", "摩擦速度，表征近地层湍流交换强度"],
            ["稳态检验", "比较全窗口与子窗口统计一致性"],
            ["ITC", "积分湍流特征，常用于湍流充分度评估"],
            ["随机不确定度", "有限湍流采样造成的统计误差"],
            ["系统误差", "校准、几何、频率响应或模型等造成的偏差"],
            ["足迹", "对当前通量贡献的上风向源区分布"],
            ["稳定度", "热力与机械湍流相对作用的状态"],
            ["Monin-Obukhov 长度", "近地层稳定度尺度 L"],
            ["粗糙度长度", "表征下垫面对风速剖面影响的尺度 z0"],
            ["零平面位移", "高冠层上方风速剖面的有效位移高度 d"],
            ["开路分析仪", "测量路径直接暴露在环境空气中"],
            ["闭路分析仪", "空气经采样管进入测量池"],
            ["Nyquist 频率", "离散采样可无混叠表示的最高频率 fs/2"],
            ["MAD", "中位数绝对偏差，稳健离散程度指标"],
            ["QC", "Quality Control，质量控制与标志"],
            ["schema", "输出字段、类型、单位和约束的结构定义"],
            ["provenance", "数据或方法的来源、参数、实现和限制"],
            ["manifest", "交付文件清单及其角色、大小和校验信息"],
        ],
        [52, 124],
        "涡动协方差与软件术语",
    )


def add_appendix_e(builder: ManualBuilder) -> None:
    builder.heading("附录 E  科学与技术参考资料", 1)
    builder.paragraph("本手册的科学说明参考以下原始论文、教材与公开技术文档，并结合 Gas EC Studio 当前代码能力重新组织。正文为独立表述，不代表外部机构对本软件的认证。")
    references = [
        "[1] Webb, E. K., Pearman, G. I., and Leuning, R. (1980). Correction of flux measurements for density effects due to heat and water vapour transfer. Quarterly Journal of the Royal Meteorological Society, 106, 85-100. DOI: 10.1002/qj.49710644707.",
        "[2] Wilczak, J. M., Oncley, S. P., and Stage, S. A. (2001). Sonic anemometer tilt correction algorithms. Boundary-Layer Meteorology, 99, 127-150. DOI: 10.1023/A:1018966204465.",
        "[3] Vickers, D. and Mahrt, L. (1997). Quality control and flux sampling problems for tower and aircraft data. Journal of Atmospheric and Oceanic Technology, 14, 512-526. DOI: 10.1175/1520-0426(1997)014<0512:QCAFSP>2.0.CO;2.",
        "[4] Foken, T. and Wichura, B. (1996). Tools for quality assessment of surface-based flux measurements. Agricultural and Forest Meteorology, 78, 83-105. DOI: 10.1016/0168-1923(95)02248-1.",
        "[5] Massman, W. J. (2000). A simple method for estimating frequency response corrections for eddy covariance systems. Agricultural and Forest Meteorology, 104, 185-198. DOI: 10.1016/S0168-1923(00)00164-7.",
        "[6] Horst, T. W. (1997). A simple formula for attenuation of eddy fluxes measured with first-order-response scalar sensors. Boundary-Layer Meteorology, 82, 219-233. DOI: 10.1023/A:1000229130034.",
        "[7] Ibrom, A. et al. (2007). Strong low-pass filtering effects on water vapour flux measurements with closed-path eddy correlation systems. Agricultural and Forest Meteorology, 147, 140-156. DOI: 10.1016/j.agrformet.2007.07.007.",
        "[8] Fratini, G. and Mauder, M. (2014). Towards a consistent eddy-covariance processing: an intercomparison of EddyPro and TK3. Atmospheric Measurement Techniques, 7, 2273-2281. DOI: 10.5194/amt-7-2273-2014.",
        "[9] Kljun, N. et al. (2015). A simple two-dimensional parameterisation for Flux Footprint Prediction (FFP). Geoscientific Model Development, 8, 3695-3713. DOI: 10.5194/gmd-8-3695-2015.",
        "[10] Kormann, R. and Meixner, F. X. (2001). An analytical footprint model for non-neutral stratification. Boundary-Layer Meteorology, 99, 207-224. DOI: 10.1023/A:1018991015119.",
        "[11] Hsieh, C. I., Katul, G., and Chi, T. (2000). An approximate analytical model for footprint estimation of scalar fluxes in thermally stratified atmospheric flows. Advances in Water Resources, 23, 765-772. DOI: 10.1016/S0309-1708(99)00042-1.",
        "[12] Mann, J. and Lenschow, D. H. (1994). Errors in airborne flux measurements. Journal of Geophysical Research, 99, 14519-14526. DOI: 10.1029/94JD00737.",
        "[13] Finkelstein, P. L. and Sims, P. F. (2001). Sampling error in eddy correlation flux measurements. Journal of Geophysical Research, 106, 3503-3509. DOI: 10.1029/2000JD900731.",
        "[14] Kaimal, J. C. and Finnigan, J. J. (1994). Atmospheric Boundary Layer Flows: Their Structure and Measurement. Oxford University Press.",
        "[15] Aubinet, M., Vesala, T., and Papale, D. (eds.) (2012). Eddy Covariance: A Practical Guide to Measurement and Data Analysis. Springer.",
        "[16] LI-COR Environmental. Eddy covariance processing software help: advanced settings, raw processing, flux calculation and output files. https://www.licor.com/support/EddyPro/topics/raw-processing-options.html (accessed 2026-07-21).",
        "[17] Gas EC Studio source documentation and method validation artifacts, software version 0.1.0 RC6.",
    ]
    for ref in references:
        p = builder.doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(ref)
        _set_run_font(r, 8.6, color=INK)
        _set_paragraph_keep(p)
    builder.callout("引用建议", "论文或正式报告引用算法时，应优先引用对应原始论文，并在方法部分记录 Gas EC Studio 软件版本、配置和实际实现路径。", "info")


def add_final_checklist(builder: ManualBuilder) -> None:
    builder.heading("交付前一页检查表", 1)
    builder.paragraph("勾选全部项目后再发布数据或报告。任何红色风险都应先闭合。")
    builder.table(
        ["检查", "内容", "完成"],
        [
            ["版本", "软件、手册、项目方法和输入批次一致", "□"],
            ["原始数据", "完整、只读留存、有清单和校验", "□"],
            ["时间", "时区、采样率、窗口边界和时钟同步正确", "□"],
            ["元数据", "站点、仪器、布局、管路和校准已核对", "□"],
            ["预检", "字段、单位、有效样本和输出目录通过", "□"],
            ["处理", "实际方法实现、回退和参数来源已检查", "□"],
            ["QC", "C 级、低置信、非稳态和低湍流已审阅", "□"],
            ["谱", "协谱、Ogive、传递函数和修正因子可解释", "□"],
            ["足迹", "来源区域和 GIS 重叠满足研究用途", "□"],
            ["不确定度", "随机误差与系统限制分开说明", "□"],
            ["报告", "结论与当前最终批次一致", "□"],
            ["证据包", "清单、方法、异常、schema 和文件均可打开", "□"],
            ["复核", "独立目录验证完成，审核人和日期已记录", "□"],
            ["归档", "原始、配置、结果、报告、日志分层备份", "□"],
        ],
        [32, 128, 16],
        "正式发布检查",
    )
    builder.callout("最后原则", "结果要能回答三件事：测到了什么、如何计算、为什么可信。无法回答其中任何一项时，就还没有完成科学交付。", "success")


def add_manual_chapters(builder: ManualBuilder, diagrams: dict[str, Path]) -> None:
    add_chapter_1(builder)
    add_chapter_2(builder, diagrams)
    add_chapter_3(builder)
    add_chapter_4(builder)
    add_chapter_5(builder)
    add_chapter_6(builder, diagrams)
    add_chapter_7(builder, diagrams)
    add_chapter_8(builder)
    add_chapter_9(builder)
    add_chapter_10(builder)
    add_chapter_11(builder)
    add_chapter_12(builder)
    add_appendix_a(builder)
    add_appendix_b(builder)
    add_appendix_c(builder)
    add_appendix_d(builder)
    add_appendix_e(builder)
    add_final_checklist(builder)


def build_manual() -> Path:
    diagrams = create_diagrams()
    builder = ManualBuilder()
    add_cover(builder)
    add_document_control(builder)
    builder.add_toc()
    # Manual chapters are appended below.
    add_manual_chapters(builder, diagrams)
    builder.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main() -> None:
    output = build_manual()
    print(output)


if __name__ == "__main__":
    main()
